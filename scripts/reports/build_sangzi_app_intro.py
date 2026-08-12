from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor


APP_NAME = "智护银龄"
DOCUMENT_TITLE = "智护银龄 APP 项目介绍"
PROJECT_ORIGIN = "华中科技大学基础医学院“慧老智治 医心为民”AI 智慧医养暑期实践项目"
TEAM_NAME = "基础医学院“慧老智治 医心为民”AI智慧医养暑期社会实践队"

BURGUNDY = "7B2431"
CHARCOAL = "20252B"
BODY_BLACK = "111111"
MUTED = "5C6168"
LINE = "C9C4BC"
LIGHT_LINE = "E4E0DA"
WARM_GRAY = "F3F0EB"
PALE_RED = "F7EFF1"
WHITE = "FFFFFF"
SUCCESS = "285B3A"
CAUTION = "7A5A00"

CONTENT_WIDTH_DXA = 8844
CONTENT_WIDTH_IN = 156 / 25.4

FONT_TITLE = "宋体"
FONT_BODY = "仿宋"
FONT_HEADING = "黑体"
FONT_SUBHEADING = "楷体"
FONT_CAPTION = "宋体"
FONT_WESTERN = "Times New Roman"


def set_explicit_fonts(r_fonts, east_asia: str, western: str) -> None:
    """Use literal font names so WPS does not substitute theme fonts differently."""
    for theme_name in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        theme_attr = qn(f"w:{theme_name}")
        if theme_attr in r_fonts.attrib:
            del r_fonts.attrib[theme_attr]
    r_fonts.set(qn("w:ascii"), western)
    r_fonts.set(qn("w:hAnsi"), western)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), western)
    r_fonts.set(qn("w:hint"), "eastAsia")


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(
    run,
    *,
    east_asia: str = FONT_BODY,
    western: str = FONT_WESTERN,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = western
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    set_explicit_fonts(r_fonts, east_asia, western)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_style_font(style, east_asia: str, western: str, size: float, color: str) -> None:
    style.font.name = western
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    set_explicit_fonts(r_fonts, east_asia, western)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("left", start),
        ("bottom", bottom),
        ("end", end),
        ("right", end),
    ):
        tag = qn(f"w:{edge}")
        node = tc_mar.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, *, color: str = LIGHT_LINE, size: int = 6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "left", "bottom", "end", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "start", "left", "bottom", "end", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = 0) -> None:
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_paragraph_bottom_border(paragraph, color: str, size: int = 6, space: int = 3) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, east_asia=FONT_CAPTION, size=10.5, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BURGUNDY)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_WESTERN)
    r_fonts.set(qn("w:hAnsi"), FONT_WESTERN)
    r_fonts.set(qn("w:eastAsia"), FONT_CAPTION)
    r_pr.extend([r_fonts, color, underline, size])
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, FONT_BODY, FONT_WESTERN, 16, BODY_BLACK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Pt(32)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(28)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.widow_control = True

    title = doc.styles["Title"]
    set_style_font(title, FONT_TITLE, FONT_WESTERN, 22, CHARCOAL)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.keep_with_next = True

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, FONT_HEADING, FONT_WESTERN, 14, MUTED)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)

    heading1 = doc.styles["Heading 1"]
    set_style_font(heading1, FONT_HEADING, FONT_WESTERN, 16, CHARCOAL)
    heading1.font.bold = True
    h1_pf = heading1.paragraph_format
    h1_pf.first_line_indent = Pt(0)
    h1_pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h1_pf.line_spacing = Pt(28)
    h1_pf.space_before = Pt(14)
    h1_pf.space_after = Pt(10)
    h1_pf.keep_with_next = True
    h1_pf.keep_together = True

    heading2 = doc.styles["Heading 2"]
    set_style_font(heading2, FONT_SUBHEADING, FONT_WESTERN, 16, CHARCOAL)
    heading2.font.bold = True
    h2_pf = heading2.paragraph_format
    h2_pf.first_line_indent = Pt(0)
    h2_pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h2_pf.line_spacing = Pt(26)
    h2_pf.space_before = Pt(10)
    h2_pf.space_after = Pt(6)
    h2_pf.keep_with_next = True
    h2_pf.keep_together = True

    caption = doc.styles["Caption"]
    set_style_font(caption, FONT_CAPTION, FONT_WESTERN, 10.5, MUTED)
    caption.font.italic = False
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_together = True

    for style_name, font_name, size, color in (
        ("Header", FONT_CAPTION, 9.5, MUTED),
        ("Footer", FONT_CAPTION, 10.5, MUTED),
    ):
        style = doc.styles[style_name]
        set_style_font(style, font_name, FONT_WESTERN, size, color)
        style.paragraph_format.first_line_indent = Pt(0)


def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(30)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(15)
    section.gutter = Mm(0)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Mm(156), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run("智护银龄 APP 项目介绍")
    set_run_font(left, east_asia=FONT_CAPTION, size=9.5, color=MUTED)
    right = p.add_run("\t正式汇报材料")
    set_run_font(right, east_asia=FONT_CAPTION, size=9.5, color=MUTED)
    set_paragraph_bottom_border(p, BURGUNDY, size=5, space=4)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.first_line_indent = Pt(0)
    dash1 = fp.add_run("— ")
    set_run_font(dash1, east_asia=FONT_CAPTION, size=10.5, color=MUTED)
    add_page_field(fp)
    dash2 = fp.add_run(" —")
    set_run_font(dash2, east_asia=FONT_CAPTION, size=10.5, color=MUTED)
    section.first_page_footer.paragraphs[0].text = ""


def add_body(doc: Document, text: str, *, first_indent: bool = True, after: float = 0) -> None:
    p = doc.add_paragraph(style="Normal")
    if not first_indent:
        p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.add_run(text)


def add_lead_body(doc: Document, lead: str, text: str, *, after: float = 2) -> None:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    lead_run = p.add_run(lead)
    set_run_font(lead_run, east_asia=FONT_HEADING, size=16, bold=True, color=CHARCOAL)
    text_run = p.add_run(text)
    set_run_font(text_run, east_asia=FONT_BODY, size=16, color=BODY_BLACK)


def add_h1(doc: Document, number: str, title: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    # 让新章节从新页开始，比在上一页末尾插入空白分页段落更稳定，
    # 可避免上一页刚好排满时产生整张空白页。
    p.paragraph_format.page_break_before = True
    n = p.add_run(f"{number}、")
    set_run_font(n, east_asia=FONT_HEADING, size=16, bold=True, color=BURGUNDY)
    t = p.add_run(title)
    set_run_font(t, east_asia=FONT_HEADING, size=16, bold=True, color=CHARCOAL)
    set_paragraph_bottom_border(p, LIGHT_LINE, size=4, space=5)


def add_h2(doc: Document, title: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(title)


def add_small_note(doc: Document, text: str, *, color: str = MUTED, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, east_asia=FONT_CAPTION, size=10.5, color=color)


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_RED)
    set_cell_borders(cell, color="D9C4C9", size=6)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(title)
    set_run_font(r1, east_asia=FONT_HEADING, size=13, bold=True, color=BURGUNDY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.first_line_indent = Pt(0)
    p2.paragraph_format.line_spacing = 1.35
    r2 = p2.add_run(text)
    set_run_font(r2, east_asia=FONT_BODY, size=14, color=BODY_BLACK)


def add_label_detail_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [1900, CONTENT_WIDTH_DXA - 1900])
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        set_cell_shading(left, WARM_GRAY)
        set_cell_borders(left)
        set_cell_borders(right)
        set_cell_margins(left, top=100, start=120, bottom=100, end=120)
        set_cell_margins(right, top=100, start=140, bottom=100, end=140)
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.space_after = Pt(0)
        lr = left.paragraphs[0].add_run(label)
        set_run_font(lr, east_asia=FONT_HEADING, size=12.5, bold=True, color=CHARCOAL)
        rr = right.paragraphs[0].add_run(value)
        set_run_font(rr, east_asia=FONT_BODY, size=13.5, color=BODY_BLACK)


def add_function_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=3)
    set_table_geometry(table, [1850, 3250, CONTENT_WIDTH_DXA - 5100])
    headers = ["功能", "用户可以做什么", "实际帮助"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BURGUNDY)
        set_cell_borders(cell, color=BURGUNDY, size=6)
        set_cell_margins(cell, top=100, start=100, bottom=100, end=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        set_run_font(run, east_asia=FONT_HEADING, size=11.5, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])

    for row_idx, values in enumerate(rows, start=1):
        for col_idx, value in enumerate(values):
            cell = table.rows[row_idx].cells[col_idx]
            if row_idx % 2 == 0:
                set_cell_shading(cell, "FAF9F7")
            set_cell_borders(cell)
            set_cell_margins(cell, top=85, start=95, bottom=85, end=95)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(
                run,
                east_asia=FONT_HEADING if col_idx == 0 else FONT_BODY,
                size=11.2,
                bold=col_idx == 0,
                color=CHARCOAL if col_idx == 0 else BODY_BLACK,
            )


def add_status_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=3)
    set_table_geometry(table, [2000, 2200, CONTENT_WIDTH_DXA - 4200])
    for idx, text in enumerate(("能力", "当前状态", "说明")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, WARM_GRAY)
        set_cell_borders(cell, color=LINE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(text)
        set_run_font(run, east_asia=FONT_HEADING, size=11.5, bold=True, color=CHARCOAL)
    set_repeat_table_header(table.rows[0])
    for ridx, values in enumerate(rows, start=1):
        for cidx, value in enumerate(values):
            cell = table.rows[ridx].cells[cidx]
            set_cell_borders(cell)
            set_cell_margins(cell, top=90, start=95, bottom=90, end=95)
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            status_color = SUCCESS if cidx == 1 and "可以" in value else BODY_BLACK
            if cidx == 1 and "完善" in value:
                status_color = CAUTION
            set_run_font(
                run,
                east_asia=FONT_HEADING if cidx < 2 else FONT_BODY,
                size=11.2,
                bold=cidx < 2,
                color=status_color,
            )


def add_image_to_paragraph(paragraph, path: Path, width_in: float, title: str, description: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    # 图片是行内对象；若沿用正文的固定行距，Word/LibreOffice 会只显示一条窄带。
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width_in))
    set_alt_text(shape, title, description)


def add_two_image_figure(
    doc: Document,
    left: tuple[Path, str, str],
    right: tuple[Path, str, str],
    *,
    width_in: float = 1.50,
) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [CONTENT_WIDTH_DXA // 2, CONTENT_WIDTH_DXA - CONTENT_WIDTH_DXA // 2])
    remove_table_borders(table)
    for cell, (path, caption, alt) in zip(table.rows[0].cells, (left, right), strict=True):
        set_cell_margins(cell, top=80, start=90, bottom=80, end=90)
        set_cell_borders(cell, color=LIGHT_LINE, size=5)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = cell.paragraphs[0]
        add_image_to_paragraph(p, path, width_in, caption, alt)
        cap = cell.add_paragraph(style="Caption")
        cap.add_run(caption)


def add_single_image_figure(
    doc: Document,
    path: Path,
    caption: str,
    alt: str,
    *,
    width_in: float = 2.25,
) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=80, start=90, bottom=80, end=90)
    p = cell.paragraphs[0]
    add_image_to_paragraph(p, path, width_in, caption, alt)
    cap = cell.add_paragraph(style="Caption")
    cap.add_run(caption)


def add_image_text_panel(
    doc: Document,
    path: Path,
    caption: str,
    alt: str,
    points: list[tuple[str, str]],
) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3200, CONTENT_WIDTH_DXA - 3200])
    remove_table_borders(table)
    image_cell, text_cell = table.rows[0].cells
    for cell in (image_cell, text_cell):
        set_cell_margins(cell, top=80, start=100, bottom=80, end=100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_borders(image_cell, color=LIGHT_LINE, size=5)
    p = image_cell.paragraphs[0]
    add_image_to_paragraph(p, path, 1.86, caption, alt)
    cap = image_cell.add_paragraph(style="Caption")
    cap.add_run(caption)
    text_cell.paragraphs[0].text = ""
    for idx, (lead, text) in enumerate(points):
        p2 = text_cell.paragraphs[0] if idx == 0 else text_cell.add_paragraph()
        p2.paragraph_format.first_line_indent = Pt(0)
        p2.paragraph_format.line_spacing = 1.35
        p2.paragraph_format.space_after = Pt(10)
        r1 = p2.add_run(lead)
        set_run_font(r1, east_asia=FONT_HEADING, size=13, bold=True, color=BURGUNDY)
        r2 = p2.add_run(text)
        set_run_font(r2, east_asia=FONT_BODY, size=13.5, color=BODY_BLACK)


def add_step_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [760, CONTENT_WIDTH_DXA - 760])
    remove_table_borders(table)
    for idx, (number, text) in enumerate(rows):
        left, right = table.rows[idx].cells
        for cell in (left, right):
            set_cell_margins(cell, top=80, start=80, bottom=80, end=80)
            set_cell_borders(cell, color=LIGHT_LINE, size=4)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(left, PALE_RED)
        p1 = left.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.first_line_indent = Pt(0)
        r1 = p1.add_run(number)
        set_run_font(r1, east_asia=FONT_HEADING, size=13, bold=True, color=BURGUNDY)
        p2 = right.paragraphs[0]
        p2.paragraph_format.first_line_indent = Pt(0)
        p2.paragraph_format.line_spacing = 1.25
        r2 = p2.add_run(text)
        set_run_font(r2, east_asia=FONT_BODY, size=13.5, color=BODY_BLACK)


def page_break(doc: Document) -> None:
    # 一级标题统一使用 page_break_before；保留调用点以维持正文结构清晰。
    return None


def build_document(output_path: Path, source_root: Path) -> None:
    screenshot_dir = source_root / "output" / "playwright"
    images = {
        "login": screenshot_dir / "report-login-current.png",
        "voice": screenshot_dir / "voice-390.png",
        "family_home": screenshot_dir / "family-home-fixed.png",
        "health": screenshot_dir / "family-health-390.png",
        "medicine": screenshot_dir / "family-medicine.png",
        "medicine_form": screenshot_dir / "medicine-plan-form.png",
        "message": screenshot_dir / "murmur-chat-390.png",
        "settings": screenshot_dir / "settings-390.png",
    }
    missing = [str(path) for path in images.values() if not path.exists()]
    if missing:
        raise FileNotFoundError("缺少 APP 运行截图：\n" + "\n".join(missing))

    doc = Document()
    setup_styles(doc)
    setup_page(doc)
    props = doc.core_properties
    props.title = DOCUMENT_TITLE
    props.subject = "智护银龄 APP 的定位、功能、使用方式、价值与边界"
    props.author = TEAM_NAME
    props.keywords = "智护银龄, 桑梓智护, 智慧医养, 适老化, APP 介绍"
    props.comments = "依据项目现行功能资料和真实运行截图制作。"
    props.created = datetime(2026, 8, 11, 12, 0, 0)
    props.modified = datetime(2026, 8, 11, 12, 0, 0)

    # 封面
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(56)
    origin = doc.add_paragraph()
    origin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    origin.paragraph_format.first_line_indent = Pt(0)
    origin.paragraph_format.space_after = Pt(16)
    run = origin.add_run(PROJECT_ORIGIN)
    set_run_font(run, east_asia=FONT_HEADING, size=13, bold=True, color=MUTED)

    line = doc.add_paragraph()
    line.paragraph_format.space_after = Pt(26)
    set_paragraph_bottom_border(line, BURGUNDY, size=10, space=0)

    title = doc.add_paragraph(style="Title")
    title.add_run("智护银龄 APP 项目介绍")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("面向长辈与家属的日常健康管理和家庭照护应用")

    statement = doc.add_paragraph()
    statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    statement.paragraph_format.first_line_indent = Pt(0)
    statement.paragraph_format.space_before = Pt(26)
    statement.paragraph_format.space_after = Pt(70)
    r = statement.add_run("正式汇报材料")
    set_run_font(r, east_asia=FONT_SUBHEADING, size=16, bold=True, color=BURGUNDY)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.first_line_indent = Pt(0)
    meta.paragraph_format.line_spacing = 1.6
    r = meta.add_run(f"{TEAM_NAME}\n当前展示版本：1.2.0\n2026年8月")
    set_run_font(r, east_asia=FONT_BODY, size=13, color=CHARCOAL)

    page_break(doc)

    # 一、APP 概况
    add_h1(doc, "一", "APP 概况")
    add_callout(
        doc,
        "一句话定位",
        "智护银龄是一款同时面向长辈和家属的智慧医养 APP：长辈可以更简单地记录健康、管理用药和使用语音助手；家属可以在获得授权后远程了解情况、协助记录，并与长辈保持日常联系。",
    )
    add_body(
        doc,
        "这款 APP 的重点不是把功能做得越来越多，而是把长辈真正需要的几件事做得更清楚、更容易操作。它把健康记录、用药提醒、家庭联系和语音陪伴放在同一个入口中，让长辈少记步骤，让家属少一点信息断层。",
    )
    add_label_detail_table(
        doc,
        [
            ("APP 名称", "智护银龄；应用界面同时使用“桑梓智护”品牌名称"),
            ("主要使用者", "需要日常健康管理的长辈，以及承担远程关怀和协助责任的家属"),
            ("核心目标", "让长辈更容易使用，让家属更清楚地了解，让双方更有边界地协作"),
            ("产品来源", "源于智慧医养实践中对适老化、家庭照护和日常健康管理需求的持续观察"),
            ("当前形态", "已经形成可实际运行的线上应用，并可在安卓手机中安装使用"),
        ],
    )
    add_small_note(doc, "说明：本页介绍的是 APP 当前定位，不代表医疗诊断或急救服务。")

    page_break(doc)

    # 二、为什么需要
    add_h1(doc, "二", "为什么需要这款 APP")
    add_body(
        doc,
        "智护银龄产生于一个朴素的判断：老年人并不是不需要数字化服务，而是很多产品并没有按照他们真实的理解方式和操作能力来设计。字号太小、入口太多、步骤太长，都会让原本有帮助的功能变得难以使用。",
    )
    add_lead_body(doc, "看不清、找不到。", "长辈常常需要在很小的文字和很多入口中寻找目标，容易犹豫，也容易误触。")
    add_lead_body(doc, "健康信息容易散。", "血压、心率、体温等数据如果只记在纸上或聊天中，时间一长就难以连续回看。")
    add_lead_body(doc, "用药容易忘，也难协助。", "长辈可能忘记时间，家属即使关心，也很难随时确认今天是否按计划服药。")
    add_lead_body(doc, "异地关怀缺少抓手。", "家属不在身边时，频繁询问容易给长辈压力；完全不问，又会担心信息不及时。")
    add_lead_body(doc, "陪伴不应越过医疗边界。", "语音助手可以帮助长辈记录、提醒和聊天，但不能替代医生给出诊断、处方或急救判断。")
    add_callout(
        doc,
        "APP 的回答",
        "把长辈端做得更简单，把家属端做得更清楚，并通过明确授权把两端连接起来。需要长辈自己决定的事情仍由长辈决定；家属只在获得相应授权后查看或协助。",
    )

    page_break(doc)

    # 三、双角色
    add_h1(doc, "三", "一款 APP，两种使用角色")
    add_body(
        doc,
        "长辈和家属面对的任务不同。智护银龄没有让所有人共用同一套复杂页面，而是根据角色分别安排入口和信息重点。这样既减少长辈的操作负担，也避免家属在自己的页面中看到不相关的信息。",
    )
    role_table = doc.add_table(rows=1, cols=2)
    set_table_geometry(role_table, [CONTENT_WIDTH_DXA // 2, CONTENT_WIDTH_DXA - CONTENT_WIDTH_DXA // 2])
    for idx, (title_text, fill, lines) in enumerate(
        [
            (
                "长辈端",
                PALE_RED,
                [
                    "首页突出大字、时间、语音入口和紧急求助",
                    "本人记录健康数据并查看最新情况",
                    "接收用药提醒并确认是否已服药",
                    "通过文字或语音与家属联系",
                    "自行决定家属可以查看或协助哪些事项",
                ],
            ),
            (
                "家属端",
                WARM_GRAY,
                [
                    "选择当前关注的长辈，查看照护看板",
                    "在获得授权后查看健康和用药情况",
                    "在获得授权后代录健康、设置用药计划",
                    "接收家庭消息和经同意分享的日常内容",
                    "管理绑定关系和分项权限",
                ],
            ),
        ]
    ):
        cell = role_table.cell(0, idx)
        set_cell_shading(cell, fill)
        set_cell_borders(cell, color=LINE)
        set_cell_margins(cell, top=160, start=170, bottom=160, end=170)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(title_text)
        set_run_font(run, east_asia=FONT_HEADING, size=15, bold=True, color=BURGUNDY if idx == 0 else CHARCOAL)
        for line_text in lines:
            p2 = cell.add_paragraph()
            p2.paragraph_format.first_line_indent = Pt(0)
            p2.paragraph_format.left_indent = Pt(6)
            p2.paragraph_format.line_spacing = 1.35
            p2.paragraph_format.space_after = Pt(6)
            r = p2.add_run(f"• {line_text}")
            set_run_font(r, east_asia=FONT_BODY, size=13, color=BODY_BLACK)
    add_small_note(doc, "角色可以在设置中切换；家属查看或代办长辈事项，需要先完成家庭绑定并获得相应授权。")

    page_break(doc)

    # 四、功能总览
    add_h1(doc, "四", "核心功能总览")
    add_function_table(
        doc,
        [
            ("智能语音", "点击麦克风说话，查看文字回应并听取播报", "减少输入和阅读负担"),
            ("健康记录", "记录并查看血压、心率、血糖、体温和体重", "把零散数据连续保存"),
            ("照护看板", "家属查看今日用药、近期趋势和异常提示", "远程了解时更有重点"),
            ("用药管理", "设置计划、接收提醒、确认服药、查看历史", "帮助按时执行日常计划"),
            ("家庭消息", "发送文字或语音，接收日常分享", "保持联系但不过度打扰"),
            ("绑定与授权", "建立家庭关系并分别控制查看、代录和管理权限", "让协助建立在同意之上"),
            ("无障碍设置", "调整字号预览和语音速度", "适应不同阅读与听取习惯"),
            ("紧急求助", "在长辈端快速发起应用内求助记录", "减少紧急时的操作步骤"),
        ],
    )
    add_body(
        doc,
        "这些功能并不是彼此分散的菜单。健康记录会进入家属获授权后的照护视图，用药确认会形成连续情况，语音助手可以帮助长辈完成部分记录，家庭消息则承担日常沟通。APP 希望把“长辈自己做”和“家属适度协助”连接起来。",
    )

    page_break(doc)

    # 五、登录与身份
    add_h1(doc, "五", "进入 APP：先确认身份，再进入对应页面")
    add_image_text_panel(
        doc,
        images["login"],
        "图 1  APP 登录页面",
        "智护银龄 APP 登录页面，包含邮箱、人机验证、验证码和登录按钮。",
        [
            ("清楚。", "页面只保留登录需要的信息，不让长辈在进入 APP 前面对过多选项。"),
            ("安全。", "通过验证问题和一次性验证码确认登录，减少密码记忆负担。"),
            ("分角色。", "首次使用时选择“长辈”或“家属”，之后进入与身份相对应的页面。"),
        ],
    )
    add_body(
        doc,
        "登录之后，APP 会根据身份自动展示不同内容。长辈端优先显示语音、用药和紧急求助；家属端优先显示照护对象、健康趋势和用药情况。角色区分不是简单换颜色，而是把不同人真正需要处理的事情放在更靠前的位置。",
    )

    page_break(doc)

    # 六、语音
    add_h1(doc, "六", "智能语音助手：让“说出来”成为主要入口")
    add_image_text_panel(
        doc,
        images["voice"],
        "图 2  智能语音助手页面",
        "智能语音助手页面，中央是醒目的麦克风按钮，顶部显示页面名称，底部提供结束对话。",
        [
            ("操作简单。", "长辈点击中央麦克风即可说话，再次点击结束本轮，不需要在多层菜单中查找。"),
            ("双重呈现。", "识别内容和回答可以显示为文字，同时也能朗读，兼顾阅读和听取习惯。"),
            ("有明确边界。", "助手可以陪伴、整理日常表达和记录完整健康数值，但不会替代医生作出诊断或更改药物。"),
        ],
    )
    add_body(
        doc,
        "当长辈说出完整、明确的健康数值时，语音助手可以在确认后帮助写入健康记录。对于日常生活内容，APP 可以先作为私密记录保存；只有长辈在当前对话中明确同意，相关内容才会分享给已绑定家属。",
    )

    page_break(doc)

    # 七、健康与看板
    add_h1(doc, "七", "健康记录与照护看板：把日常情况看清楚")
    add_two_image_figure(
        doc,
        (
            images["family_home"],
            "图 3  家属照护看板（演示数据）",
            "家属照护看板，展示今日用药、七日服药完成情况、血压、心率和近期趋势；姓名和数值为演示数据。",
        ),
        (
            images["health"],
            "图 4  长辈健康页面（演示数据）",
            "长辈健康页面，展示心率、血压、血糖、体温和体重；姓名和数值为演示数据。",
        ),
    )
    add_h2(doc, "长辈看到什么")
    add_body(
        doc,
        "长辈可以看到自己的最新心率、血压、血糖、体温和体重，并继续录入新的测量结果。页面按指标分开显示，数值、单位和时间同时出现，避免只看到一个孤立数字。",
    )
    add_h2(doc, "家属看到什么")
    add_body(
        doc,
        "家属先选择正在照护的长辈，再查看今日用药、近七日服药完成情况、血压、心率和异常提示。没有相应授权时，APP 会明确显示“尚未授权”，而不是把没有权限的数据误显示为零。",
    )
    add_small_note(doc, "以上页面用于功能说明，姓名和所有健康数值均为演示数据，不代表真实个人健康状况。")

    page_break(doc)

    # 八、用药
    add_h1(doc, "八", "用药管理：从提醒到确认，再到家属协助")
    add_two_image_figure(
        doc,
        (
            images["medicine"],
            "图 5  当日用药与计划管理（演示数据）",
            "用药页面，显示当日计划和现有用药计划；姓名、药品和时间为演示数据。",
        ),
        (
            images["medicine_form"],
            "图 6  添加用药计划页面（演示数据）",
            "添加用药计划页面，包含药品名称、剂量、时间和起止日期；内容为演示数据。",
        ),
    )
    add_body(
        doc,
        "用药管理不是简单地弹出一个闹钟。APP 先保存药品名称、剂量、服药时间和日期范围，再按计划提醒长辈。长辈可以确认已服、延后或跳过；这些结果会保留在历史中，便于自己回看，也便于获授权家属了解近期执行情况。",
    )
    add_body(
        doc,
        "家属在获得用药管理授权后，可以为长辈新增、调整或停用计划，也可以协助确认。APP 会把“计划是什么”和“实际完成情况”分开保存，避免后来修改计划时把过去情况一并改掉。",
    )
    add_callout(
        doc,
        "安全提示",
        "APP 只负责记录和提醒，不判断药物是否适合，也不建议剂量。药品名称、剂量和服用方式必须以医生或药师给出的方案为准。",
    )

    page_break(doc)

    # 九、家庭联系
    add_h1(doc, "九", "家庭联系：沟通、绑定和授权放在一起")
    add_two_image_figure(
        doc,
        (
            images["message"],
            "图 7  家庭消息与日常分享（演示数据）",
            "家庭对话页面，展示一条经同意分享的日常内容；文字和时间为演示内容。",
        ),
        (
            images["settings"],
            "图 8  设置与绑定管理（演示数据）",
            "设置页面，包含个人信息、绑定管理、无障碍设置、消息通知和角色切换；姓名为演示身份。",
        ),
    )
    add_h2(doc, "先绑定，再协作")
    add_body(
        doc,
        "长辈可以生成短时有效的绑定码，家属输入绑定码并选择与长辈的关系后，双方建立家庭联系。绑定成功并不等于家属可以查看全部内容；健康查看、健康代录、用药管理和紧急通知分别由长辈授权。",
    )
    add_h2(doc, "日常消息不只是一句问候")
    add_body(
        doc,
        "双方可以发送文字或语音消息。语音助手整理出的日常内容默认只属于长辈本人；只有长辈明确同意后，APP 才会把这段内容作为“碎碎念”分享给已绑定家属，让陪伴建立在尊重和同意之上。",
    )
    add_small_note(doc, "本页画面为功能演示，姓名、消息内容和时间均为演示信息。")

    page_break(doc)

    # 十、适老化
    add_h1(doc, "十", "适老化与无障碍：不是把字放大这么简单")
    add_lead_body(doc, "看得清。", "正文、按钮和关键数值使用更大的字号与更高的对比度，重要操作具有清楚边界。")
    add_lead_body(doc, "点得准。", "常用按钮留出更大的触控范围，减少按钮拥挤和误触。")
    add_lead_body(doc, "找得到。", "长辈端首页只突出少量高频入口，返回、确认和取消的位置保持清楚。")
    add_lead_body(doc, "听得懂。", "语音回答尽量简短、口语化；用户可以调整语音速度，也可以通过文字同步确认。")
    add_lead_body(doc, "失败时有说明。", "录音、识别、保存或网络出现问题时，页面给出明确反馈，不让用户停在没有解释的空白状态。")
    add_body(
        doc,
        "设置页面提供字号预览和语音速度调整。这里的目标不是为所有长辈规定同一种“大字模式”，而是让不同视力、听力和操作习惯的人能够选择更适合自己的方式。",
    )
    add_callout(
        doc,
        "设计原则",
        "让长辈先理解，再操作；先给出最常用的选择，再提供更多设置；重要结果同时用文字和声音说明。",
    )

    page_break(doc)

    # 十一、紧急求助
    add_h1(doc, "十一", "紧急求助：减少步骤，但不夸大能力")
    add_body(
        doc,
        "长辈端首页设置醒目的紧急求助按钮。点击后，APP 会发起一条紧急求助记录，并在页面中反馈是否成功。这样做的价值是让长辈在慌乱时少找一个入口、少完成几步操作。",
    )
    add_h2(doc, "当前已经可以做到")
    add_body(
        doc,
        "长辈可以在 APP 内直接发起求助；系统会保存求助记录并识别需要通知的家庭联系人；页面会提示提交结果，失败时会提醒用户立即拨打 120。",
    )
    add_h2(doc, "当前仍不能替代")
    add_body(
        doc,
        "电话、短信或系统推送等外部通知渠道仍需进一步接入和真实验证。因此，目前的紧急求助功能不能替代 120，也不能承诺家属一定会在手机系统层面即时收到通知。",
    )
    add_callout(
        doc,
        "必须记住",
        "遇到胸痛、呼吸困难、意识不清、疑似卒中、严重外伤等紧急情况，应立即拨打 120 或当地急救电话，不应等待 APP 回复。",
    )

    page_break(doc)

    # 十二、典型用法
    add_h1(doc, "十二", "典型使用方式")
    add_h2(doc, "场景一：长辈独立完成一天的日常管理")
    add_step_table(
        doc,
        [
            ("1", "早晨打开 APP，首页以大字显示时间和主要入口；到服药时间后收到提醒。"),
            ("2", "服药后点击确认，APP 保存本次完成情况；如果暂时不能服用，可以选择延后。"),
            ("3", "测量血压或心率后，手动输入，或对语音助手说出完整数值并确认记录。"),
            ("4", "想与家人联系时发送文字或语音；想分享一天中的小事时，由本人决定是否发送给家属。"),
        ],
    )
    add_h2(doc, "场景二：家属在异地提供适度协助")
    add_step_table(
        doc,
        [
            ("1", "家属打开照护看板，先选择当前关注的长辈。"),
            ("2", "在授权范围内查看今日用药是否完成、近期健康数值和需要关注的异常提示。"),
            ("3", "需要时为长辈调整用药计划或代录健康数据；没有授权的项目保持锁定。"),
            ("4", "通过消息了解长辈近况，用具体问题代替反复追问，让关心更有依据。"),
        ],
    )

    page_break(doc)

    # 十三、价值
    add_h1(doc, "十三", "APP 的核心价值")
    add_lead_body(doc, "对长辈：降低使用门槛。", "把常用功能放在更显眼的位置，用大字、语音和清楚反馈减少记忆步骤，让长辈保留更多自主操作的可能。")
    add_lead_body(doc, "对家属：把担心变成可理解的信息。", "在获得授权后，家属看到的是明确的健康记录、用药完成情况和日常消息，而不是依靠猜测判断长辈近况。")
    add_lead_body(doc, "对家庭：建立有边界的协作。", "绑定关系和分项授权让“关心”与“隐私”同时被看见，家属可以帮忙，但不能默认获得全部信息。")
    add_lead_body(doc, "对智慧医养服务：形成连续的日常记录。", "APP 不替代医疗服务，而是补充医院和家庭之间的日常管理环节，使健康、用药和沟通信息更连续。")
    add_lead_body(doc, "对适老化产品探索：验证以使用者为中心的设计。", "APP 把语音优先、双角色、明确反馈和权限控制落实到实际页面，为后续真实试用和改进提供可观察的产品基础。")
    add_callout(
        doc,
        "价值判断",
        "智护银龄的价值不在于替老人和家属作决定，而在于让长辈更容易表达和记录，让家属在得到同意后更及时地理解和协助。",
    )

    page_break(doc)

    # 十四、完成度
    add_h1(doc, "十四", "当前完成度与后续完善")
    add_status_table(
        doc,
        [
            ("身份与角色", "已经可以使用", "登录、首次选择身份、角色切换和个人资料已形成完整使用路径。"),
            ("健康记录", "已经可以使用", "长辈可记录和查看；家属在授权后可查看或代录。"),
            ("用药管理", "已经可以使用", "支持计划、提醒、确认、历史和家属协助。"),
            ("家庭联系", "已经可以使用", "支持绑定、分项授权、文字和语音消息。"),
            ("智能语音", "已经可以使用", "支持说话、识别、文字回答和朗读，并可帮助记录完整健康数值。"),
            ("无障碍设置", "已经可以使用", "支持字号预览和语音速度调整。"),
            ("紧急求助", "仍需继续完善", "应用内发起和记录已可用，外部电话、短信或推送通知仍需接入和验证。"),
            ("真实使用验证", "仍需继续完善", "需要更多长辈和家属在真实设备、真实网络和连续使用场景中反馈。"),
        ],
    )
    add_body(
        doc,
        "下一步工作的重点不是继续增加大量功能，而是把已经具备的功能做得更可靠：完善外部紧急通知，验证提醒是否真正到达，优化首次使用指导，持续检查隐私授权，并根据真实长辈和家属的反馈调整页面。",
    )

    page_break(doc)

    # 十五、边界和资料
    add_h1(doc, "十五", "使用边界与资料说明")
    add_h2(doc, "医疗与安全边界")
    add_body(
        doc,
        "智护银龄用于日常健康记录、用药提醒、家庭联系和语音陪伴。它不提供疾病诊断，不开具处方，不判断药物剂量，不替代医生、药师或护理人员，也不替代 120 等急救渠道。任何异常健康数值都应结合本人症状和医生意见判断。",
    )
    add_h2(doc, "隐私与演示说明")
    add_body(
        doc,
        "家属查看或代办长辈事项需要家庭绑定和相应授权。本文全部 APP 截图来自真实运行页面，但截图中的姓名、健康数值、药品、消息和时间均为功能演示，不代表真实个人信息或服务效果。",
    )
    add_h2(doc, "文档编排依据")
    sources = [
        (
            "GB/T 148—1997《印刷、书写和绘图纸幅面尺寸》",
            "https://openstd.samr.gov.cn/bzgk/std/newGbInfo?hcno=20746CFEE63514B24DD64A415CB65377",
        ),
        (
            "GB/T 9704—2012《党政机关公文格式》",
            "https://openstd.samr.gov.cn/bzgk/std/newGbInfo?hcno=F3CC9BEF482524C895FDA7A08BB4A70E",
        ),
        (
            "GB/T 15834—2011《标点符号用法》",
            "https://openstd.samr.gov.cn/bzgk/std/newGbInfo?hcno=22EA6D162E4110E752259661E1A0D0A8",
        ),
        (
            "GB/T 15835—2011《出版物上数字用法》",
            "https://openstd.samr.gov.cn/bzgk/std/newGbInfo?hcno=F5DAC3377DA99C8D78AE66735B6359C7",
        ),
    ]
    for idx, (label, url) in enumerate(sources, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"{idx}. ")
        set_run_font(r, east_asia=FONT_CAPTION, size=10.5, color=MUTED)
        add_hyperlink(p, label, url)

    add_small_note(
        doc,
        "项目内容依据：智护银龄现行功能说明、1.2.0 版本说明、关于我们页面及 2026 年 8 月 13 日线上运行页面。",
    )
    add_callout(
        doc,
        "结语",
        "智护银龄希望把健康管理和家庭关怀放回日常生活：长辈能够更容易地使用，家属能够在获得同意后更清楚地协助，重要信息能够被连续地记录和理解。",
    )
    # WPS 会在以表格结束的文档后补一个默认空段落，极端情况下会产生空白尾页。
    # 显式保留一个最小终止段落，确保它稳定停留在结语之后的同一页。
    terminal = doc.add_paragraph()
    terminal.paragraph_format.first_line_indent = Pt(0)
    terminal.paragraph_format.space_before = Pt(0)
    terminal.paragraph_format.space_after = Pt(0)
    terminal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    terminal.paragraph_format.line_spacing = Pt(1)
    set_run_font(terminal.add_run(""), east_asia=FONT_CAPTION, size=1, color=WHITE)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Generated: {output_path}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="生成智护银龄 APP 正式介绍 Word 文档")
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("docs/reports/智护银龄APP项目介绍.docx"),
        help="输出 DOCX 路径",
    )
    parser.add_argument(
        "--source-root",
        type=Path,
        default=Path(__file__).resolve().parents[2],
        help="包含 output/playwright 运行截图的项目根目录",
    )
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build_document(args.output.resolve(), args.source_root.resolve())
